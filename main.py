import json
from pathlib import Path
from starlette.background import BackgroundTasks
import os
from fastapi import FastAPI, Form,File,UploadFile,Response,status
from fastapi.responses import FileResponse
import datetime
from docxtpl import DocxTemplate
import textwrap
import re 
from docx.shared import Inches

today = datetime.datetime.today()
today_in_one_week =  today + datetime.timedelta(days=7)
doc = DocxTemplate("./form2revised.docx")

my_app = FastAPI()
def remove_file(path: str) -> None:
    os.unlink(path)

@my_app.post("/getInformation",status_code=200)
async def handle_form( response: Response,background_tasks: BackgroundTasks,name:str=Form(),fathers_name_or_husbands_name:str=Form(),signature: UploadFile = File(...),dob:str=Form(),gender:str=Form(),maritalstatus:str=Form(),pf_number:str=Form(),address:str=Form(),epf_nominee_details:str=Form(),eps_member_details:str=Form(),eps_nominee:str=Form()):
 try:
        FILE_TYPES = set(['png', 'jpeg'])
        image = f"{signature.filename}" 
        if '.' in image and image.rsplit('.', 1)[1] in FILE_TYPES:   
          with open(image, "wb+") as file_object:
                file_object.write(signature.file.read())
        else:
                raise FileNotFoundError
      
        epf=json.loads(epf_nominee_details)
        eps_members=json.loads(eps_member_details)
        eps=json.loads(eps_nominee)
        context = {
                "name": name,
                "fathers_name_or_husbands_name":fathers_name_or_husbands_name,
                "dob":dob,
                "gender": gender,
                "maritalstatus":maritalstatus,
                "pf_number": pf_number,
                "address":textwrap.fill(address,width=240),
                "TODAY": today.strftime("%d-%m-%Y"), 
                "TODAY_IN_ONE_WEEK": today_in_one_week.strftime("%d-%m-%Y"),
        }
        for i in range(len(epf)):
                context[f"name_and_address_of_nominee{i+1}"]=epf[i]["name_and_address_of_nominee"]
                context[f"nominee{i+1}_relationship"]=epf[i]["nominee_relationship"]
                context[f"dob{i+1}"]=epf[i]["dob"]
                context[f"totalamt_or_share{i+1}"]=epf[i]["totalamt_or_share"]
                context[f"if_nominee{i+1}_is_a_minor_mention_guardian_name_and_address"]=epf[i]["if_nominee_is_a_minor_mention_guardian_name_and_address"]
        for i in range(len(eps_members)):
                context[f"name_address_of_the_family_member{i+1}"]=eps_members[i]["name_address_of_the_family_member"]
                context[f"epsdob{i+1}"]=eps_members[i]["epsdob"]
                context[f"relationship{i+1}_with_member"]=eps_members[i]["relationship_with_member"]
        for i in range(len(epf)):
                context[f"name_and_address_of_nominee{i+1}_eps"]= eps[i]["name_and_address_of_nominee_eps"]
                context[f"dob{i+1}epsnominee"]= eps[i]["dobepsnominee"]
                context[f"relation{i+1}"]= eps[i]["relation"]
        

        doc.render(context)
        
        img_tag = re.compile(r'%') # declare pattern
        #STEP 1 
        for _p in enumerate(doc.paragraphs): 

                img_paragraph = None
                if bool(img_tag.match(_p[1].text)): 
                        
                        img_paragraph = _p[1] 
                #STEP 2
                temp_text = img_tag.split(img_paragraph.text)
                print(temp_text)
                img_paragraph.runs[0].text = temp_text[0]
                _r = img_paragraph.add_run()
                _r.add_picture(image, width = Inches(1.25))
                img_paragraph.add_run(temp_text[1])
        

 except FileNotFoundError : 
        response.status_code = status.HTTP_415_UNSUPPORTED_MEDIA_TYPE
        return {        
        "error" : "FILE TYPE IS NOT PNG/JPEG"
                  } 
 except :
        response.status_code = status.HTTP_500_INTERNAL_SERVER_ERROR
        return {
        "error" : "docx file not rendered"
                  }
 else:
        doc.save(f"{context['name']}-form2.docx")
        document_path = Path(__file__).parent /f"{context['name']}-form2.docx"
        os.remove(image) 
        background_tasks.add_task(remove_file, document_path)  
        return FileResponse(path=document_path,filename=f"{context['name']}-form2.docx")

 
